import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
import os
from wordcloud import WordCloud
import docx
from docx.shared import Inches
from io import BytesIO

st.set_page_config(
    page_title="DataSensei",
    page_icon="😎",
    initial_sidebar_state="collapsed",
)

if os.path.exists("header.html"):
    with open("header.html", "r") as file:
        header_html_content = file.read()
    st.markdown(header_html_content, unsafe_allow_html=True)

st.title("DataSensei - Visualization Studio")

uploaded_file = st.file_uploader("Upload a CSV or XLSX file", type=["csv", "xlsx"])

if uploaded_file is not None:
    filename = uploaded_file.name.lower()
    if filename.endswith(".csv"):
        df = pd.read_csv(uploaded_file)
    else:
        df = pd.read_excel(uploaded_file)

    st.subheader("Data Preview")
    st.write(df.head())

    available_plots = [
        "Line Plot",
        "Scatter Plot",
        "Bar Plot (Vertical and Horizontal)",
        "Grouped Bar Plot",
        "Histogram",
        "Box Plot",
        "Violin Plot",
        "Density Plot",
        "Heatmap",
        "Word Cloud",
    ]

    image_dir = "images"
    os.makedirs(image_dir, exist_ok=True)

    selected_plots = st.multiselect("Select plots to visualize", available_plots)

    image_paths = []


    for plot_type in selected_plots:
        if plot_type == "Line Plot":
            x_column = st.selectbox(
                "Select X-axis column for Line Plot", df.columns, key="line_x"
            )
            y_column = st.selectbox(
                "Select Y-axis column for Line Plot", df.columns, key="line_y"
            )

            fig, ax = plt.subplots(figsize=(8, 3))
            ax.plot(df[x_column], df[y_column])
            ax.set_xlabel(x_column)
            ax.set_ylabel(y_column)
            ax.set_title("Line Plot")
            st.pyplot(fig)

            line_plot_image_path = os.path.join(image_dir, "line_plot.png")
            fig.savefig(line_plot_image_path, bbox_inches="tight")
            plt.close(fig)
            image_paths.append(line_plot_image_path)

        elif plot_type == "Scatter Plot":
            x_column = st.selectbox(
                "Select X-axis column for Scatter Plot", df.columns, key="scatter_x"
            )
            y_column = st.selectbox(
                "Select Y-axis column for Scatter Plot", df.columns, key="scatter_y"
            )

            fig, ax = plt.subplots(figsize=(5, 3))
            ax.scatter(df[x_column], df[y_column])
            ax.set_xlabel(x_column)
            ax.set_ylabel(y_column)
            ax.set_title("Scatter Plot")
            st.pyplot(fig)

            scatter_plot_image_path = os.path.join(image_dir, "scatter_plot.png")
            fig.savefig(scatter_plot_image_path, bbox_inches="tight")
            plt.close(fig)
            image_paths.append(scatter_plot_image_path)

        elif plot_type == "Bar Plot (Vertical and Horizontal)":
            x_column = st.selectbox(
                "Select X-axis column for Bar Plot", df.columns, key="bar_x"
            )
            y_column = st.selectbox(
                "Select Y-axis column for Bar Plot", df.columns, key="bar_y"
            )
            orientation = st.radio(
                "Select orientation", ["vertical", "horizontal"], key="bar_orientation"
            )

            fig, ax = plt.subplots(figsize=(5, 3))
            if orientation == "vertical":
                ax.bar(df[x_column], df[y_column])
                ax.set_xlabel(x_column)
                ax.set_ylabel(y_column)
            else:
                ax.barh(df[x_column], df[y_column])
                ax.set_xlabel(y_column)
                ax.set_ylabel(x_column)
            ax.set_title("Bar Plot")
            for label in ax.get_xticklabels():
                label.set_rotation(45)
            st.pyplot(fig)

            bar_plot_image_path = os.path.join(image_dir, "bar_plot.png")
            fig.savefig(bar_plot_image_path, bbox_inches="tight")
            plt.close(fig)
            image_paths.append(bar_plot_image_path)

        elif plot_type == "Grouped Bar Plot":
            x_column = st.selectbox(
                "Select X-axis column for Grouped Bar Plot",
                df.columns,
                key="grouped_bar_x",
            )
            y_columns = st.multiselect(
                "Select Y-axis columns for Grouped Bar Plot",
                df.columns,
                key="grouped_bar_y",
            )

            if len(y_columns) < 1:
                st.warning("Please select at least one Y-axis column for Grouped Bar Plot.")
            else:
                grouped_bar_df = df[y_columns].copy()
                grouped_bar_df[x_column] = df[x_column]
                grouped_bar_df.set_index(x_column, inplace=True)

                fig, ax = plt.subplots(figsize=(6, 3))
                grouped_bar_df.plot(kind="bar", ax=ax)
                ax.set_xlabel(x_column)
                ax.set_ylabel("Value")
                ax.set_title("Grouped Bar Plot")
                for label in ax.get_xticklabels():
                    label.set_rotation(45)
                st.pyplot(fig)

                grouped_bar_plot_image_path = os.path.join(
                    image_dir, "grouped_bar_plot.png"
                )
                fig.savefig(grouped_bar_plot_image_path, bbox_inches="tight")
                plt.close(fig)
                image_paths.append(grouped_bar_plot_image_path)

        elif plot_type == "Histogram":
            numeric_cols = df.select_dtypes(include="number").columns
            if len(numeric_cols) == 0:
                st.warning("No numeric columns available for Histogram.")
                continue

            column = st.selectbox(
                "Select column for Histogram", numeric_cols, key="hist_col"
            )
            fig, ax = plt.subplots(figsize=(5, 3))
            ax.hist(df[column].dropna(), bins=20, edgecolor="k")
            ax.set_xlabel(column)
            ax.set_ylabel("Frequency")
            ax.set_title("Histogram")
            st.pyplot(fig)

            hist_image_path = os.path.join(image_dir, "histogram.png")
            fig.savefig(hist_image_path, bbox_inches="tight")
            plt.close(fig)
            image_paths.append(hist_image_path)

        elif plot_type == "Box Plot":
            numeric_cols = df.select_dtypes(include="number").columns
            if len(numeric_cols) == 0:
                st.warning("No numeric columns available for Box Plot.")
                continue

            column = st.selectbox(
                "Select column for Box Plot", numeric_cols, key="box_col"
            )
            fig, ax = plt.subplots(figsize=(5, 2.5))
            ax.boxplot(df[column].dropna(), vert=False)
            ax.set_xlabel(column)
            ax.set_title("Box Plot")
            st.pyplot(fig)

            box_plot_image_path = os.path.join(image_dir, "box_plot.png")
            fig.savefig(box_plot_image_path, bbox_inches="tight")
            plt.close(fig)
            image_paths.append(box_plot_image_path)

        elif plot_type == "Violin Plot":
            numeric_cols = df.select_dtypes(include="number").columns
            if len(numeric_cols) == 0:
                st.warning("No numeric columns available for Violin Plot.")
                continue

            column = st.selectbox(
                "Select column for Violin Plot", numeric_cols, key="violin_col"
            )
            fig, ax = plt.subplots(figsize=(5, 3))
            sns.violinplot(data=df, y=column, ax=ax)
            ax.set_ylabel(column)
            ax.set_title("Violin Plot")
            st.pyplot(fig)

            violin_plot_image_path = os.path.join(image_dir, "violin_plot.png")
            fig.savefig(violin_plot_image_path, bbox_inches="tight")
            plt.close(fig)
            image_paths.append(violin_plot_image_path)

        elif plot_type == "Density Plot":
            numeric_cols = df.select_dtypes(include="number").columns
            if len(numeric_cols) == 0:
                st.warning("No numeric columns available for Density Plot.")
                continue

            column = st.selectbox(
                "Select column for Density Plot", numeric_cols, key="density_col"
            )
            fig, ax = plt.subplots(figsize=(5, 3))
            sns.kdeplot(df[column].dropna(), shade=True, ax=ax)
            ax.set_xlabel(column)
            ax.set_ylabel("Density")
            ax.set_title("Density Plot")
            st.pyplot(fig)

            density_plot_image_path = os.path.join(image_dir, "density_plot.png")
            fig.savefig(density_plot_image_path, bbox_inches="tight")
            plt.close(fig)
            image_paths.append(density_plot_image_path)

        elif plot_type == "Heatmap":
            numeric_df = df.select_dtypes(include="number")

            if numeric_df.shape[1] < 2:
                st.warning("Need at least two numeric columns for a Heatmap.")
                continue

            corr_matrix = numeric_df.corr()

            fig, ax = plt.subplots(figsize=(6, 4))
            sns.heatmap(corr_matrix, annot=True, cmap="coolwarm", fmt=".2f", ax=ax)
            ax.set_title("Correlation Heatmap")
            st.pyplot(fig)

            heatmap_image_path = os.path.join(image_dir, "heatmap.png")
            fig.savefig(heatmap_image_path, bbox_inches="tight")
            plt.close(fig)
            image_paths.append(heatmap_image_path)

        elif plot_type == "Word Cloud":
            text_cols = df.select_dtypes(include="object").columns
            if len(text_cols) == 0:
                st.warning("No text columns available for Word Cloud.")
                continue

            text_column = st.selectbox(
                "Select Text column for Word Cloud", text_cols, key="wordcloud_col"
            )
            text_series = df[text_column].dropna().astype(str)

            if text_series.empty:
                st.warning("Selected text column is empty.")
                continue

            wordcloud = WordCloud(
                width=800, height=400, background_color="white"
            ).generate(" ".join(text_series))

            fig, ax = plt.subplots(figsize=(6, 3))
            ax.imshow(wordcloud, interpolation="bilinear")
            ax.axis("off")
            ax.set_title("Word Cloud")
            st.pyplot(fig)

            wordcloud_image_path = os.path.join(image_dir, "wordcloud.png")
            fig.savefig(wordcloud_image_path, bbox_inches="tight")
            plt.close(fig)
            image_paths.append(wordcloud_image_path)


    if selected_plots and image_paths:
        if st.button("Generate and Save Graphs to Word"):
            doc = docx.Document()

            for plot_type, image_path in zip(selected_plots, image_paths):
                st.image(image_path, caption=plot_type, use_column_width=True)

                doc.add_heading(plot_type, level=1)
                doc.add_picture(image_path, width=Inches(6))

            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            st.download_button(
                label="Download graphs.docx",
                data=buffer,
                file_name="graphs.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )

            st.success("Graphs have been added to Word and are ready to download.")
    else:
        if selected_plots:
            st.info("Plots will appear here once generated.")
else:
    st.info("Please upload a CSV or XLSX file to begin.")
