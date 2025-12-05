import streamlit as st
import pandas as pd
from docx import Document
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
from io import BytesIO
from sklearn.linear_model import LinearRegression, LogisticRegression
from docx.shared import Inches
from sklearn.cluster import KMeans
from sklearn.model_selection import train_test_split
from sklearn.metrics import accuracy_score
from sklearn.cluster import AgglomerativeClustering
from sklearn.preprocessing import StandardScaler, LabelEncoder
from sklearn.neighbors import KNeighborsClassifier
from scipy.cluster.hierarchy import dendrogram, linkage
import os

st.set_page_config(
    page_title="DataSensei", page_icon="😎", initial_sidebar_state="collapsed"
)

# Optional header HTML
if os.path.exists("header.html"):
    with open("header.html", "r") as header_file:
        ht = header_file.read()
        st.markdown(ht, unsafe_allow_html=True)

doc = Document()
doc.add_heading("Data Analytics Report", 0)

def add_plot_to_doc(plt_figure_bytes, doc, width_inches=6):
    """Helper to add a BytesIO plot to the Word doc safely."""
    plt_figure_bytes.seek(0)
    doc.add_picture(plt_figure_bytes, width=Inches(width_inches))


def get_classification_data(data):
    """Common feature/target selection for classification models."""
    st.write("### Select Features and Target for Classification")

    feature_columns = st.multiselect(
        "Select numerical feature columns",
        data.select_dtypes(include=[np.int64, np.float64]).columns,
    )

    target_column = st.selectbox(
        "Select target column (classification)",
        data.columns,
    )

    if not feature_columns or not target_column:
        st.warning("Please select at least one feature and a target column.")
        return None, None, None, None, None

    X = data[feature_columns].dropna()
    y = data.loc[X.index, target_column]

    if y.dtype == object:
        le = LabelEncoder()
        y = le.fit_transform(y)
    else:
        y = y.astype(int)

    if len(np.unique(y)) < 2:
        st.error("Target column must have at least two classes for classification.")
        return None, None, None, None, None

    X_train, X_test, y_train, y_test = train_test_split(
        X, y, test_size=0.3, random_state=42
    )

    return X_train, X_test, y_train, y_test, feature_columns

def perform_anomaly_detection(data, doc):
    st.subheader("Anomaly Detection")

    doc.add_paragraph("Performing Anomaly Detection...")

    selected_columns = st.multiselect(
        "Select numerical columns for Anomaly Detection",
        data.select_dtypes(include=[np.int64, np.float64]).columns,
    )

    if len(selected_columns) > 0:
        st.success("Report generated successfully!")

        doc.add_heading("Anomaly Detection", level=2)

        for col in selected_columns:

            if col in data.columns:
                column_data = data[col]

                if column_data.isnull().values.any():
                    msg = f"Column '{col}' contains missing values. Please clean or impute before anomaly detection."
                    st.warning(msg)
                    doc.add_paragraph(msg)
                    continue

                st.markdown(f"###  Boxplot for `{col}`")

                from scipy.stats import zscore
                z_scores = zscore(column_data)
                threshold = 3
                anomalies = column_data[abs(z_scores) > threshold]

                fig, ax = plt.subplots(figsize=(10, 5))
                sns.boxplot(x=column_data, ax=ax, color="skyblue")
                sns.stripplot(x=anomalies, ax=ax, color="red", size=8, label='Anomalies')

                ax.set_title(f"Boxplot with Anomalies: {col}")
                ax.set_xlabel(col)
                ax.legend()

                st.pyplot(fig)

                plot_bytes = BytesIO()
                fig.savefig(plot_bytes, format='png', bbox_inches="tight")
                plot_bytes.seek(0)
                plt.close(fig)

                doc.add_heading(f"Anomaly Detection for {col}", level=3)
                add_plot_to_doc(plot_bytes, doc)

                st.write(f"**Anomaly Threshold (Z-Score):** {threshold}")
                st.write(f"**Detected Anomalies Count:** {len(anomalies)}")

                if len(anomalies) > 0:
                    st.write("Detected anomaly values:")
                    st.dataframe(anomalies.to_frame(name=col))
                else:
                    st.success("No anomalies detected!")

                doc.add_paragraph(f"Anomalies detected in '{col}': {len(anomalies)}")
                doc.add_paragraph(f"Z-Score Threshold: {threshold}")
                doc.add_page_break()




def perform_customer_segmentation(data, doc):
    st.subheader("Customer Segmentation (Agglomerative Clustering)")

    doc.add_paragraph("Performing Customer Segmentation and Analysis...")

    selected_columns = st.multiselect(
        "Select numerical columns for Customer Segmentation",
        data.select_dtypes(include=[np.int64, np.float64]).columns,
    )

    if len(selected_columns) > 1:
        st.success("Segmentation successfully completed!")

        doc.add_heading("Customer Segmentation", level=2)
        doc.add_heading(
            f"Segmentation Using Columns: {', '.join(selected_columns)}",
            level=3
        )

        X = data[selected_columns].dropna()
        X_std = StandardScaler().fit_transform(X)

        n_clusters = 3
        clustering = AgglomerativeClustering(n_clusters=n_clusters)
        cluster_labels = clustering.fit_predict(X_std)

        fig_scatter, ax_scatter = plt.subplots(figsize=(10, 6))
        scatter = ax_scatter.scatter(
            X_std[:, 0],
            X_std[:, 1],
            c=cluster_labels,
            cmap="viridis"
        )
        ax_scatter.set_xlabel("Dimension 1")
        ax_scatter.set_ylabel("Dimension 2")
        ax_scatter.set_title("Customer Segmentation (Agglomerative Clustering)")

        st.markdown("### Cluster Scatter Plot")
        st.pyplot(fig_scatter)

        scatter_bytes = BytesIO()
        fig_scatter.savefig(scatter_bytes, format="png", bbox_inches="tight")
        scatter_bytes.seek(0)
        plt.close(fig_scatter)

        add_plot_to_doc(scatter_bytes, doc)

        Z = linkage(X_std, method="ward")

        fig_dendro, ax_dendro = plt.subplots(figsize=(12, 6))
        dendrogram(Z, ax=ax_dendro)
        ax_dendro.set_title("Dendrogram (Hierarchical Clustering)")
        ax_dendro.set_xlabel("Samples")
        ax_dendro.set_ylabel("Distance")

        st.markdown("### Dendrogram")
        st.pyplot(fig_dendro)

        dendro_bytes = BytesIO()
        fig_dendro.savefig(dendro_bytes, format="png", bbox_inches="tight")
        dendro_bytes.seek(0)
        plt.close(fig_dendro)

        add_plot_to_doc(dendro_bytes, doc)

        counts = (
            pd.Series(cluster_labels, name="Cluster")
            .value_counts()
            .sort_index()
            .reset_index()
            .rename(columns={"index": "Cluster", "Cluster": "Count"})
        )

        st.markdown("### 📊 Cluster Analysis Summary")
        st.dataframe(counts)

        doc.add_heading("Cluster Analysis Results", level=3)
        doc.add_paragraph(f"Number of Clusters: {n_clusters}")
        doc.add_paragraph("Cluster sizes:")
        doc.add_paragraph(counts.to_string(index=False))

        doc.add_page_break()
    else:
        st.warning("Please select at least two numerical columns!")




def perform_descriptive_statistics(data, doc):
    st.subheader("Descriptive Statistics")

    doc.add_paragraph("Performing Descriptive Statistics...")

    selected_columns = st.multiselect(
        "Select numerical columns for Descriptive Statistics",
        data.select_dtypes(include=[np.int64, np.float64]).columns,
    )

    if len(selected_columns) > 0:
        st.success("Report generated successfully!")

        doc.add_heading("Descriptive Statistics", level=2)

        for col in selected_columns:
            doc.add_heading(f"Descriptive Statistics for Column: {col}", level=3)

            if col in data.columns:
                desc_stats = data[col].describe()
                st.write(f"**{col}**")
                st.write(desc_stats)
                doc.add_paragraph(desc_stats.to_string())

                doc.add_page_break()


def regression_analysis(data, doc):
    st.subheader("Multiple Linear Regression Analysis")

    doc.add_paragraph("Performing Multiple Linear Regression Analysis...")

    numeric_cols = data.select_dtypes(include=[np.number]).columns

    if len(numeric_cols) < 2:
        st.warning("Need at least 2 numeric columns for multiple linear regression.")
        return

    predictor_cols = st.multiselect(
        "Select predictor variables (independent variables)",
        numeric_cols,
        key="mlr_predictors"
    )

    target_col = st.selectbox(
        "Select target variable (dependent variable)",
        [c for c in numeric_cols if c not in predictor_cols],
        key="mlr_target"
    )

    st.markdown("###  Model Settings")
    col1, col2 = st.columns(2)
    with col1:
        test_size = st.slider("Test Size (fraction)", 0.1, 0.5, 0.2, 0.05)
    with col2:
        random_state = st.number_input("Random State", min_value=0, max_value=9999, value=42)

    run_button = st.button("Run Regression")

    if run_button and predictor_cols and target_col:
        # Prepare data
        model_df = data[predictor_cols + [target_col]].dropna()
        if model_df.empty:
            st.error("Not enough valid data after removing missing values.")
            return

        X = model_df[predictor_cols]
        y = model_df[target_col]

        X_train, X_test, y_train, y_test = train_test_split(
            X, y,
            test_size=test_size,
            random_state=random_state
        )

        model = LinearRegression()
        model.fit(X_train, y_train)

        y_pred = model.predict(X_test)

        tolerance = 0.10
        accuracy = np.mean(np.abs((y_pred - y_test) / y_test) <= tolerance)

        coef_df = pd.DataFrame({
            "Predictor": predictor_cols,
            "Coefficient": model.coef_
        })

        st.success("Regression successfully executed!")

        st.markdown("### 📌 Model Summary")
        st.write(f"**Target:** `{target_col}`")
        st.write(f"**Predictors:** {', '.join(predictor_cols)}")
        st.write(f"**Accuracy :** `{accuracy*100:.2f}%`")

        st.markdown("### 📊 Coefficients")
        st.dataframe(coef_df)

        st.markdown("### 🔹Intercept Value")
        st.write(model.intercept_)

        sample_results = pd.DataFrame({
            "Actual": y_test.values,
            "Predicted": y_pred
        }).head(20)
        st.markdown("### 🔸Sample Test Predictions")
        st.dataframe(sample_results)

        doc.add_heading("Multiple Linear Regression Analysis", level=2)
        doc.add_paragraph(f"Target: {target_col}")
        doc.add_paragraph(f"Predictors: {', '.join(predictor_cols)}")
        doc.add_paragraph(f"Accuracy : {accuracy*100:.2f}%")

        doc.add_heading("Coefficients", level=3)
        doc.add_paragraph(coef_df.to_string(index=False))

        doc.add_heading("Intercept", level=3)
        doc.add_paragraph(str(model.intercept_))

        doc.add_page_break()




def logistic_regression_classification(data, doc):
    st.subheader("Logistic Regression Classification")
    doc.add_heading("Logistic Regression Classification", level=2)

    X_train, X_test, y_train, y_test, feature_columns = get_classification_data(data)
    if X_train is None:
        return

    model = LogisticRegression(max_iter=1000)
    model.fit(X_train, y_train)
    y_pred = model.predict(X_test)

    acc = accuracy_score(y_test, y_pred)

    st.write(f"Logistic Regression Accuracy: {acc:.2f}")
    doc.add_paragraph(f"Logistic Regression Accuracy: {acc:.2f}")
    doc.add_page_break()


def knn_classification(data, doc):
    st.subheader("K-Nearest Neighbors (KNN) Classification")
    doc.add_heading("K-Nearest Neighbors (KNN) Classification", level=2)

    X_train, X_test, y_train, y_test, feature_columns = get_classification_data(data)
    if X_train is None:
        return

    k = st.slider("Select number of neighbors (k)", min_value=1, max_value=15, value=5)

    model = KNeighborsClassifier(n_neighbors=k)
    model.fit(X_train, y_train)
    y_pred = model.predict(X_test)

    acc = accuracy_score(y_test, y_pred)

    st.write(f"KNN Accuracy (k={k}): {acc:.2f}")
    doc.add_paragraph(f"KNN Accuracy (k={k}): {acc:.2f}")
    doc.add_page_break()


def kmeans_clustering(data, doc):
    st.subheader("K-Means Clustering")
    doc.add_heading("K-Means Clustering", level=2)

    feature_columns = st.multiselect(
        "Select numerical columns for K-Means clustering",
        data.select_dtypes(include=[np.int64, np.float64]).columns,
    )

    if len(feature_columns) < 2:
        st.warning("Please select at least two numerical columns for clustering.")
        return

    n_clusters = st.slider("Select number of clusters", min_value=2, max_value=10, value=3)

    X = data[feature_columns].dropna()

    kmeans = KMeans(n_clusters=n_clusters, random_state=42)
    cluster_labels = kmeans.fit_predict(X)

    st.write("Cluster centers:")
    centers_df = pd.DataFrame(kmeans.cluster_centers_, columns=feature_columns)
    st.dataframe(centers_df)

    doc.add_paragraph(f"K-Means Clustering with {n_clusters} clusters.")
    doc.add_paragraph("Cluster Centers:")
    doc.add_paragraph(centers_df.to_string(index=False))

    fig, ax = plt.subplots(figsize=(10, 6))
    scatter = ax.scatter(X.iloc[:, 0], X.iloc[:, 1], c=cluster_labels)
    ax.set_xlabel(feature_columns[0])
    ax.set_ylabel(feature_columns[1])
    ax.set_title("K-Means Clustering")

    st.pyplot(fig)

    clustering_plot = BytesIO()
    fig.savefig(clustering_plot, format="png", bbox_inches="tight")
    clustering_plot.seek(0)
    plt.close(fig)

    add_plot_to_doc(clustering_plot, doc)
    doc.add_page_break()



def main():
    st.title("DataSensei")
    st.write("Please upload your data in CSV format.")

    uploaded_file = st.file_uploader("Upload CSV file", type=["csv"])

    if uploaded_file is not None:
        data = pd.read_csv(uploaded_file)

        st.write("### Data Preview:")
        st.write(data.head())

        doc.add_paragraph("Data Preview:")
        doc.add_paragraph(data.head().to_string())

        st.write("### Data Info:")
        st.write(f"Number of Rows: {data.shape[0]}")
        st.write(f"Number of Columns: {data.shape[1]}")

        st.success("Data uploaded successfully!")

        st.subheader("Choose an Analytics Task")

        analytics_task = st.selectbox(
            "Select an Analytics Task",
            [
                "Descriptive Statistics",
                "Anomaly Detection",
                "Linear Regression",
                "Logistic Regression",
                "KNN Classification",
                "KMeans Clustering",
                "Customer Segmentation",
            ],
        )

        if analytics_task == "Anomaly Detection":
            perform_anomaly_detection(data, doc)

        elif analytics_task == "Customer Segmentation":
            perform_customer_segmentation(data, doc)

        elif analytics_task == "Descriptive Statistics":
            perform_descriptive_statistics(data, doc)

        elif analytics_task == "Linear Regression":
            regression_analysis(data, doc)

        elif analytics_task == "KMeans Clustering":
            kmeans_clustering(data, doc)

        elif analytics_task == "Logistic Regression":
            logistic_regression_classification(data, doc)

        elif analytics_task == "KNN Classification":
            knn_classification(data, doc)

        else:
            st.warning("Please select at least one analysis technique.")
        report_buffer = BytesIO()
        doc.save(report_buffer)
        report_buffer.seek(0)

        st.download_button(
            label="Download Data Analytics Report",
            data=report_buffer,
            file_name="data_analytics_report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )


if __name__ == "__main__":
    main()

